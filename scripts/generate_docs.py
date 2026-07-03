"""Generate DOCX from markdown research documents with embedded images."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR = Path(r"c:\Users\agaev\Documents\GitHub\llm-generation-design\docs\research\carbon-fiber-net-defense")

IMAGE_MAP = {
    "cf-net-defense-concept.png": BASE_DIR / "cf-net-defense-concept.png",
    "cost-comparison-infographic.png": BASE_DIR / "cost-comparison-infographic.png",
    "border-deployment-map.png": BASE_DIR / "border-deployment-map.png",
    "altitude-threat-coverage.png": BASE_DIR / "altitude-threat-coverage.png",
    "net-curtain-concept.png": BASE_DIR / "net-curtain-concept.png",
    "diagram-defense-layers.png": BASE_DIR / "diagram-defense-layers.png",
    "diagram-material-weight.png": BASE_DIR / "diagram-material-weight.png",
    "diagram-lebanon-deployment.png": BASE_DIR / "diagram-lebanon-deployment.png",
    "diagram-golan-deployment.png": BASE_DIR / "diagram-golan-deployment.png",
    "diagram-jordan-deployment.png": BASE_DIR / "diagram-jordan-deployment.png",
    "diagram-roadmap.png": BASE_DIR / "diagram-roadmap.png",
    "diagram-curtain-concept.png": BASE_DIR / "diagram-curtain-concept.png",
    "diagram-self-protection.png": BASE_DIR / "diagram-self-protection.png",
}


def md_to_docx(md_path: Path, docx_path: Path, title: str, lang: str = "en"):
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    content = md_path.read_text(encoding='utf-8')
    lines = content.split('\n')

    i = 0
    in_table = False
    table_headers = []
    table_rows = []
    in_code_block = False
    code_lines = []
    code_lang = ""

    while i < len(lines):
        line = lines[i]

        if line.startswith('```'):
            if in_code_block:
                if code_lang == 'mermaid':
                    p = doc.add_paragraph()
                    p.style = doc.styles['Normal']
                    run = p.add_run('[Mermaid Diagram - view in Markdown renderer]\n' + '\n'.join(code_lines))
                    run.font.name = 'Consolas'
                    run.font.size = Pt(8)
                    run.font.italic = True
                else:
                    p = doc.add_paragraph()
                    p.style = doc.styles['Normal']
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                code_lines = []
                in_code_block = False
                code_lang = ""
            else:
                in_code_block = True
                code_lang = line[3:].strip()
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        img_match = re.match(r'!\[.*?\]\((.+?)\)', line)
        if img_match:
            img_file = img_match.group(1)
            img_path = IMAGE_MAP.get(img_file)
            if img_path and img_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(img_path), width=Inches(6.0))
            else:
                p = doc.add_paragraph(f'[Image: {img_file}]')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if '|' in line and line.count('|') >= 2 and not line.strip().startswith('```'):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells and all(c.strip() == '' or set(c.strip()) <= {'-', ':', ' '} for c in cells):
                i += 1
                continue
            if not in_table:
                in_table = True
                table_headers = cells
                table_rows = []
            else:
                table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            _flush_table(doc, table_headers, table_rows)
            in_table = False
            table_headers = []
            table_rows = []

        if line.startswith('# ') and not line.startswith('## '):
            heading = doc.add_heading(line[2:].strip(), level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=3)
        elif line.startswith('> '):
            text = line[2:].strip()
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            p = doc.add_paragraph(text)
            try:
                p.style = doc.styles['Intense Quote']
            except KeyError:
                pass
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            doc.add_paragraph(text, style='List Bullet')
        elif line.startswith('---'):
            doc.add_page_break()
        elif line.strip() == '':
            pass
        else:
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
            clean = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', clean)
            clean = re.sub(r'`(.+?)`', r'\1', clean)
            if clean.strip():
                doc.add_paragraph(clean.strip())

        i += 1

    if in_table and table_headers and table_rows:
        _flush_table(doc, table_headers, table_rows)

    doc.save(str(docx_path))
    print(f"Generated: {docx_path}")


def _flush_table(doc, headers, rows):
    if not headers:
        return
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(headers[:num_cols]):
        cell = table.rows[0].cells[j]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    for row_idx, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data[:num_cols]):
            cell = table.rows[row_idx + 1].cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()


if __name__ == '__main__':
    for name in ["research-en", "executive-summary-en", "executive-summary-he"]:
        md = BASE_DIR / f"{name}.md"
        docx = BASE_DIR / f"{name}.docx"
        if md.exists():
            md_to_docx(md, docx, name)

    he_md = BASE_DIR / "research-he.md"
    he_docx = BASE_DIR / "research-he.docx"
    if he_md.exists():
        md_to_docx(he_md, he_docx, "research-he", lang="he")

    print("\nAll DOCX files generated successfully.")
